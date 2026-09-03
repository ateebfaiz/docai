"""DocAI full pipeline: OCR cascade, document parsing, entity extraction, cleaning, pretty output."""
import csv
import io
import json
import re
import sqlite3
from datetime import datetime, date
from pathlib import Path
from typing import Any, Optional

import numpy as np
import pandas as pd
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

# ---------------------------------------------------------------------------
# Engine availability probe
# ---------------------------------------------------------------------------

_AVAIL: dict[str, bool] = {}


def _try_import(name: str) -> bool:
    if name not in _AVAIL:
        try:
            __import__(name)
            _AVAIL[name] = True
        except Exception:
            _AVAIL[name] = False
    return _AVAIL[name]


# ---------------------------------------------------------------------------
# 1. IMAGE PREPROCESSING — every cleaning trick available
# ---------------------------------------------------------------------------

def preprocess_image(path: str | Path, variants: bool = True) -> list[Image.Image]:
    """Return [original, enhanced] or a full variant set for OCR confidence."""
    img = Image.open(path).convert("RGB")
    out = [img]
    if variants:
        # grayscale + contrast boost (helps most OCR)
        g = ImageOps.grayscale(img)
        g = ImageEnhance.Contrast(g).enhance(2.0)
        out.append(g)
        # upscaled 2x (helps small text)
        big = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
        big = ImageOps.grayscale(big)
        big = ImageEnhance.Sharpness(big).enhance(1.5)
        out.append(big)
        # binarized (threshold) for crisp text
        bw = ImageOps.grayscale(img).point(lambda p: 255 if p > 140 else 0)
        out.append(bw)
        # denoised
        d = ImageOps.grayscale(img).filter(ImageFilter.MedianFilter(3))
        out.append(d)
    return out


# ---------------------------------------------------------------------------
# 2. OCR CASCADE — docling (structure) → paddleocr → easyocr → rapidocr → tesseract
# ---------------------------------------------------------------------------

_converter = None
_ocr_engines: dict[str, Any] = {}


def _docling():
    global _converter
    if _converter is None and _try_import("docling.document_converter"):
        from docling.document_converter import DocumentConverter
        _converter = DocumentConverter()
    return _converter


def _engine_paddle():
    if "paddle" not in _ocr_engines:
        if _try_import("paddleocr"):
            try:
                from paddleocr import PaddleOCR
                _ocr_engines["paddle"] = PaddleOCR(use_angle_cls=True, lang="en", show_log=False)
            except Exception as e:
                print(f"paddleocr init failed: {e}")
                _ocr_engines["paddle"] = None
        else:
            _ocr_engines["paddle"] = None
    return _ocr_engines.get("paddle")


def _engine_easy():
    if "easy" not in _ocr_engines:
        if _try_import("easyocr"):
            try:
                import easyocr
                _ocr_engines["easy"] = easyocr.Reader(["en"], gpu=False, verbose=False)
            except Exception as e:
                print(f"easyocr init failed: {e}")
                _ocr_engines["easy"] = None
        else:
            _ocr_engines["easy"] = None
    return _ocr_engines.get("easy")


def _engine_rapid():
    if "rapid" not in _ocr_engines:
        if _try_import("rapidocr"):
            try:
                from rapidocr import RapidOCR
                _ocr_engines["rapid"] = RapidOCR()
            except Exception as e:
                print(f"rapidocr init failed: {e}")
                _ocr_engines["rapid"] = None
        else:
            _ocr_engines["rapid"] = None
    return _ocr_engines.get("rapid")


def _engine_tesseract():
    if "tess" not in _ocr_engines:
        if _try_import("pytesseract"):
            try:
                import pytesseract
                _ocr_engines["tess"] = pytesseract
            except Exception as e:
                print(f"tesseract init failed: {e}")
                _ocr_engines["tess"] = None
        else:
            _ocr_engines["tess"] = None
    return _ocr_engines.get("tess")


def ocr_image(path: str | Path) -> tuple[str, dict]:
    """OCR a single image through every available engine; return (best_text, engine_report)."""
    text = ""
    report: dict[str, Any] = {}
    variants = preprocess_image(path)
    engines = {
        "paddleocr": _engine_paddle(),
        "easyocr": _engine_easy(),
        "rapidocr": _engine_rapid(),
        "tesseract": _engine_tesseract(),
    }
    for name, eng in engines.items():
        if eng is None:
            report[name] = "unavailable"
            continue
        try:
            eng_text = ""
            for v in variants:
                img_bytes = io.BytesIO()
                v.save(img_bytes, format="PNG")
                img_bytes.seek(0)
                try:
                    res = eng(img_bytes)
                except Exception:
                    # some engines prefer a path or numpy array
                    v2 = v.convert("RGB")
                    arr = np.array(v2)
                    res = eng(arr)
                if res is None:
                    continue
                # normalize per-engine output
                lines: list[str] = []
                if name == "paddleocr":
                    for page in res:
                        for item in page:
                            lines.append(item[1][0] if isinstance(item, (list, tuple)) and len(item) > 1 and item[1] else "")
                elif name == "easyocr":
                    for item in res:
                        lines.append(item[1] if isinstance(item, (list, tuple)) and len(item) > 1 else "")
                elif name == "rapidocr":
                    if isinstance(res, tuple):
                        res, _ = res
                    for item in res or []:
                        lines.append(item[1] if len(item) > 1 else "")
                elif name == "tesseract":
                    lines = res.splitlines()
                piece = "\n".join(l for l in lines if l)
                if len(piece) > len(eng_text):
                    eng_text = piece
            report[name] = f"{len(eng_text)} chars"
            if len(eng_text) > len(text):
                text = eng_text
        except Exception as e:
            report[name] = f"error: {e}"
    return text, report


# ---------------------------------------------------------------------------
# 3. DOCUMENT PARSING — docling returns structure + tables + markdown
# ---------------------------------------------------------------------------

def parse_document(path: str | Path) -> dict[str, Any]:
    """Docling: PDF/image/docx/xlsx/html → structure + tables + md.
    Fast path: disable table-structure model and internal OCR (major CPU hogs).
    Our OCR cascade handles image regions; pdfplumber-style text extraction handles text layers.
    """
    result: dict[str, Any] = {"markdown": "", "tables": [], "text": "", "error": None, "page_count": 0}
    conv = _docling()
    if conv is None:
        result["error"] = "docling unavailable"
        return result
    try:
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        opts = PdfPipelineOptions()
        opts.do_ocr = False                  # our cascade handles OCR when needed
        opts.do_table_structure = False      # biggest CPU hog; tables come from text layer
        opts.generate_page_images = False
        opts.generate_picture_images = False
        conv.pipeline_options = opts
        doc = conv.convert(str(path), pipeline_options=opts).document
        result["markdown"] = doc.export_to_markdown()
        result["text"] = doc.export_to_text()
        tables = []
        for tbl in doc.tables:
            try:
                df = tbl.export_to_dataframe()
                tables.append(df.to_dict(orient="records"))
            except Exception:
                continue
        result["tables"] = tables
        result["page_count"] = len(getattr(doc, "pages", []) or [])
    except Exception as e:
        result["error"] = str(e)
    return result


# ---------------------------------------------------------------------------
# 4. ENTITY EXTRACTION — money, dates, phones, emails, invoice numbers, ids
# ---------------------------------------------------------------------------

_MONEY_RE = re.compile(r"(?:US\$|USD|\$|PKR|Rs\.?)\s?([0-9][0-9,]*(?:\.[0-9]{2})?)", re.I)
_DATE_RE = re.compile(
    r"\b(\d{1,4})[/\-.](\d{1,2})[/\-.](\d{2,4})\b"
    r"|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+(\d{1,2}),?\s+(\d{4})\b",
    re.I,
)
_PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}|\+?9\d{9,}")
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_INVOICE_RE = re.compile(r"(?:invoice|inv|bill)\s*(?:no|#|number)?\.?\s*:?\s*([A-Za-z0-9\-/]{4,})", re.I)
_CNIC_RE = re.compile(r"\b\d{5}[\s\-]?\d{7}[\s\-]?\d\b")
_ACCOUNT_RE = re.compile(r"(?:account|a/c|iban)\s*(?:no|#|number)?\.?:?\s*([A-Z0-9]{6,})", re.I)
_ORDER_RE = re.compile(r"(?:order|po)\s*(?:no|#|number|id)?\.?\s*:?\s*([A-Za-z0-9\-]{4,})", re.I)
_REGISTRATION_RE = re.compile(r"(?:registration|reg)\s*(?:no|#|number)?\.?\s*:?\s*([A-Za-z0-9\-]{4,})", re.I)
_TAXYEAR_RE = re.compile(r"\btax\s*(?:year|period)\s*:\s*(\d{4})", re.I)
_NAME_RE = re.compile(r"(?:name)\s*:\s*([A-Z][A-Z \.]+)", re.I)


# Entity validation guards — reject false positives from OCR noise & split words.
_STOPWORDS = {"address", "contact", "mobile", "phone", "name", "no", "number",
              "date", "due", "status", "system", "medium", "period", "registration",
              "description", "code", "amount", "total", "subtotal", "grand",
              "individual", "company", "invoice", "receipt", "bank", "account"}
_RETURN_OR_FIELD = {"name", "address", "contact", "total", "taxable", "income",
                    "registration", "status", "due", "date", "medium", "system",
                    "period", "amount", "description"}


def _valid_id(s: str) -> bool:
    """Reject obvious OCR truncations and stop-words for ID fields."""
    v = s.strip()
    if len(v) < 4:
        return False
    # Must be alphanumeric with at least one digit (real IDs/CINCs/orders)
    if not any(ch.isdigit() for ch in v):
        return False
    # Reject bare stopwords (e.g. "istration", "status", "address")
    if v.lower() in _STOPWORDS or v.lower().lstrip(".") in _RETURN_OR_FIELD:
        return False
    # Reject truncated fragments (single letter prefix, or fragments like "istration")
    if v[0] in "aeiouAEIOU" and len(v) < 6 and not v.isdigit():
        return False
    return True


def _valid_name(s: str) -> bool:
    """Reject label/header words captured as 'names' (address:, contact:, etc.)."""
    v = s.strip()
    if len(v) < 3:
        return False
    if v.lower().rstrip(".") in _RETURN_OR_FIELD or v.lower() in _STOPWORDS:
        return False
    # A name normally has ≥2 words or is a single recognizable name (not "Address"/"Status")
    if len(v.split()) == 1 and not v.istitle():
        return False
    return True


def extract_entities(text: str) -> dict[str, list[str]]:
    """Extract structured fields from raw text via regex layers + validation."""
    return {
        "amounts": list(dict.fromkeys(_MONEY_RE.findall(text))),
        "dates": list(dict.fromkeys(_DATE_RE.findall(text))),
        "phones": list(dict.fromkeys(_PHONE_RE.findall(text))),
        "emails": list(dict.fromkeys(_EMAIL_RE.findall(text))),
        "invoice_numbers": list(dict.fromkeys(
            m for m in _INVOICE_RE.findall(text) if m.strip() and _valid_id(m))),
        "cnic_ids": list(dict.fromkeys(_CNIC_RE.findall(text))),
        "account_numbers": list(dict.fromkeys(
            m for m in _ACCOUNT_RE.findall(text) if _valid_id(m))),
        "order_numbers": list(dict.fromkeys(
            m for m in _ORDER_RE.findall(text) if _valid_id(m))),
        "registration_numbers": list(dict.fromkeys(
            m for m in _REGISTRATION_RE.findall(text) if _valid_id(m))),
        "names": list(dict.fromkeys(
            m.strip() for m in _NAME_RE.findall(text) if _valid_name(m))),
    }


# ---------------------------------------------------------------------------
# 5. ADAPTIVE, CONTENT-AWARE CLASSIFICATION
# ---------------------------------------------------------------------------

# Strong positive signals (weighted by how distinctive they are)
_TAXONOMY: dict[str, list[tuple[str, float]]] = {
    "invoice": [
        ("invoice", 3), ("total due", 3), ("amount payable", 3),
        ("balance due", 2), ("bill to", 2), ("payment terms", 1), ("due date", 1),
    ],
    "tax": [
        ("federal board of revenue", 5), ("fbr", 4), ("tax", 2), ("tax year", 3),
        ("income tax", 3), ("self assessment", 4), ("return", 1), ("taxable income", 3),
        ("registration no", 2), ("fil source", 1), ("tax reduction", 2),
    ],
    "bank_statement": [
        ("account statement", 4), ("bank statement", 5), ("transaction", 2),
        ("account summary", 3), ("debit", 2), ("credit", 2), ("opening balance", 3),
        ("closing balance", 3), ("bank", 1),
    ],
    "employment": [
        ("salary", 2), ("payroll", 3), ("employer", 2), ("employee", 2),
        ("offer letter", 3), ("wages", 2), ("pay slip", 4), ("remuneration", 3),
    ],
    "immigration": [
        ("passport", 4), ("visa", 4), ("immigration", 4), ("travel document", 3),
        ("residence permit", 3), ("border", 1),
    ],
    "receipt": [
        ("receipt", 3), ("payment received", 3), ("thank you for your", 2),
        ("paid", 1), ("total paid", 2),
    ],
    "identity": [
        ("cnic", 4), ("national identity", 4), ("driving license", 3),
        ("license no", 2), ("nic", 3), ("identity card", 3),
    ],
    "order": [
        ("order id", 3), ("order number", 3), ("order no", 3), ("shipment", 2),
        ("delivery", 2), ("shipping", 1), ("tracking", 1), ("customer", 1),
    ],
    "legal": [
        ("court", 3), ("complaint", 3), ("affidavit", 3), ("legal notice", 3),
        ("plaintiff", 3), ("defendant", 3), ("judgment", 2), ("petition", 2),
    ],
    "personal": [
        ("profile", 2), ("personal information", 2), ("address", 1),
        ("contact", 1), ("mobile", 1), ("subscriber", 1),
    ],
    "metadata": [
        ("manifest", 4), ("master guide", 4), ("reorganization", 4),
        ("execution plan", 3), ("metadata", 3), ("legacy", 2),
        ("documentation", 2), ("guide", 2), ("readme", 2),
    ],
}

# Minimum confidence to trust an automated classification; below this we quarantine.
_MIN_CONFIDENCE = 0.50

# Negative signals that down-rank a category (content-aware: avoid false positives)
_NEGATIVE: dict[str, list[str]] = {
    "order": ["order to make", "tax order"],       # FBR "Order" doc is tax, not order
    "invoice": ["tax invoice", "order for"],       # tax invoice ≠ invoice
    "personal": ["registration", "self assessment"],
    "tax": ["master guide", "manifest", "reorganization", "execution plan", "metadata"],
    "identity": ["manifest", "master guide", "reorganization", "execution plan", "metadata"],
}
_NEGATIVE_WEIGHT = -2.0

# Entity → category hint (adaptive: extracted content informs classification)
_ENTITY_HINTS: dict[str, str] = {
    "invoice_numbers": "invoice",
    "cnic_ids": "identity",
    "registration_numbers": "tax",
    "order_numbers": "order",
    "account_numbers": "bank_statement",
}


def _flatten(entity_values: dict) -> str:
    """Flatten extracted entities into a searchable string."""
    parts = []
    for k, v in entity_values.items():
        if isinstance(v, list):
            parts.extend(str(x) for x in v)
    return " ".join(parts).lower()


def classify(text: str, entities: dict | None = None) -> dict[str, Any]:
    """Content-aware classification using weighted keywords + entity hints + negatives."""
    low = text.lower()
    scores: dict[str, float] = {}

    # Weighted positive keyword matches
    for cat, kws in _TAXONOMY.items():
        score = 0.0
        matched = []
        for kw, weight in kws:
            if kw in low:
                score += weight
                matched.append(kw)
        if score:
            scores[cat] = score

    # Adaptive entity hints (boost based on what was actually extracted)
    if entities:
        flat = _flatten(entities)
        for ent_key, ent_cat in _ENTITY_HINTS.items():
            v = entities.get(ent_key)
            if isinstance(v, list) and v:
                # Only boost if the entity is meaningful (exclude pure digit noise)
                meaningful = [x for x in v if any(c.isalpha() for c in str(x)) or len(str(x)) > 5]
                if meaningful:
                    scores[ent_cat] = scores.get(ent_cat, 0) + 2.0

    # Apply negative signals (down-rank false positives)
    for cat, negs in _NEGATIVE.items():
        if any(neg in low for neg in negs):
            scores[cat] = scores.get(cat, 0) + _NEGATIVE_WEIGHT

    top = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    if not top:
        return {"category": "uncategorized", "confidence": 0.0, "matches": {}, "used_hints": False}

    # Confidence = top score / total positive score (exclude negatives turned zero)
    positive_total = sum(v for v in scores.values() if v > 0)
    if positive_total <= 0:
        return {"category": "uncategorized", "confidence": 0.0, "matches": {},
                "used_hints": bool(entities)}

    winner, winner_score = top[0]
    # If winner got dragged to <=0 by negatives, fall back to highest positive
    if winner_score <= 0:
        winner = max(scores, key=lambda k: scores[k])
        winner_score = scores[winner]
        if winner_score <= 0:
            return {"category": "uncategorized", "confidence": 0.0, "matches": {},
                    "used_hints": bool(entities)}

    return {
        "category": winner,
        "confidence": round(winner_score / positive_total, 2),
        "matches": {cat: round(s, 2) for cat, s in scores.items() if s > 0},
        "used_hints": bool(entities),
    }


# ---------------------------------------------------------------------------
# 6. ADAPTIVE PER-DOCUMENT FIELD EXTRACTION (content-specific schema)
# ---------------------------------------------------------------------------

# For each doc type, define which fields are relevant. Only include non-empty ones.
_TYPE_FIELDS: dict[str, list[tuple[str, str]]] = {
    "invoice": [
        ("invoice_number", "invoice_numbers"),
        ("total_amount", "amounts"),
        ("due_date", "dates"),
        ("vendor", "names"),
    ],
    "tax": [
        ("registration_no", "registration_numbers"),
        ("tax_year", "tax_year"),
        ("taxable_income", "amounts"),
        ("taxpayer", "names"),
    ],
    "bank_statement": [
        ("account_no", "account_numbers"),
        ("transactions", "dates"),
        ("amounts", "amounts"),
    ],
    "order": [
        ("order_no", "order_numbers"),
        ("amounts", "amounts"),
        ("customer", "names"),
        ("dates", "dates"),
    ],
    "identity": [
        ("cnic_no", "cnic_ids"),
        ("names", "names"),
    ],
    "receipt": [
        ("receipt_amount", "amounts"),
        ("date", "dates"),
    ],
    "employment": [
        ("salary", "amounts"),
        ("employee", "names"),
    ],
    "personal": [
        ("names", "names"),
        ("phones", "phones"),
        ("emails", "emails"),
        ("cnic_ids", "cnic_ids"),
    ],
}


def _extract_tax_year(text: str) -> str:
    # Strict: "tax year:" or "tax period:" followed by a 4-digit year
    m = _TAXYEAR_RE.search(text)
    if m:
        return m.group(1)
    # Relaxed: 4-digit year that appears as its own token right after "year"/"period"
    for pat in (r"\byear[:\s]*(20\d{2})\b", r"\bperiod[:\s]*(20\d{2})\b"):
        mm = re.search(pat, text, re.I)
        if mm:
            return mm.group(1)
    # Range fallback (e.g. "01-Jul-2024 - 30-Jun-2025"): prefer the END year as the tax year
    m_range = re.search(r"(20\d{2})\s*[-–]\s*(?:[\dA-Za-z\- ]+)?(20\d{2})", text)
    if m_range:
        return m_range.group(2)
    # Last resort: a bare 20xx year that isn't part of a date range — take the most common
    years = re.findall(r"\b(20\d\d)\b", text)
    if not years:
        return ""
    from collections import Counter
    year = Counter(years).most_common(1)[0][0]
    return year


def _get_entities_for_field(text: str, field_src: str, entities: dict) -> list[str]:
    """Pull the appropriate entity list or inject a computed field."""
    if field_src in entities:
        return [str(x) for x in entities[field_src] if x]
    if field_src == "tax_year":
        y = _extract_tax_year(text)
        return [y] if y else []
    return []


def extract_typed_fields(text: str, doc_type: str, entities: dict) -> dict[str, list[str]]:
    """Content-adaptive: extract only fields relevant to the detected doc type."""
    if doc_type not in _TYPE_FIELDS:
        # Fallback: generic extraction for unknown types
        return {k: [str(x) for x in v if x] for k, v in entities.items()}
    result: dict[str, list[str]] = {}
    for field_name, field_src in _TYPE_FIELDS[doc_type]:
        values = _get_entities_for_field(text, field_src, entities)
        if values:
            result[field_name] = values
    return result


# ---------------------------------------------------------------------------
# 7. CLEANING — normalize whitespace, unicode, OCR garbage
# ---------------------------------------------------------------------------

def clean_text(raw: str) -> str:
    """Remove OCR artifacts and normalize whitespace/encoding."""
    raw = raw.replace("\u2014", "-").replace("\u2013", "-")
    raw = raw.replace("\u201c", '"').replace("\u201d", '"')
    raw = raw.replace("\u2018", "'").replace("\u2019", "'")
    raw = "".join(ch for ch in raw if ch >= " " or ch in "\n\t")
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in raw.splitlines()]
    return "\n".join(l for l in lines if l)


# ---------------------------------------------------------------------------
# 8. PRETTY OUTPUT — structured report with rich metadata
# ---------------------------------------------------------------------------

def make_report(source: str, parse: dict, entities: dict, cls: dict, typed_fields: dict,
                ocr_report: dict | None = None, engine: str = "") -> dict:
    """Assemble a clean, human-readable extraction report."""
    now = datetime.now().isoformat()
    return {
        "source": str(source),
        "processed_at": now,
        "engine": engine or ("docling" if not parse.get("error") else "ocr-cascade"),
        "document_type": cls.get("category", "uncategorized"),
        "classification_confidence": cls.get("confidence", 0),
        "classification_matches": cls.get("matches", {}),
        "entities": {k: v for k, v in entities.items() if v},
        "fields": typed_fields,                       # content-adaptive typed fields
        "tables": parse.get("tables", []),
        "tables_found": len(parse.get("tables", [])),
        "page_count": parse.get("page_count", 0),
        "chars_extracted": len(parse.get("text", "") or ""),
        "ocr_engine_report": ocr_report or {},
        "cleaned_text": clean_text(parse.get("text", "") or ""),
        "markdown": parse.get("markdown", "")[:20000],
    }


def run_full_pipeline(path: str | Path) -> dict:
    """Run every tool: preprocess → parse (docling) → OCR fallback → entities → adaptive classify."""
    p = Path(path)
    ext = p.suffix.lower()
    _IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp", ".heic"}
    _TEXT_EXTS = {".txt", ".csv", ".md", ".json", ".xml", ".html", ".htm", ".log"}

    # Native text formats: read directly, NEVER through image OCR.
    if ext in _TEXT_EXTS:
        text = _read_fallback(p)
        parse: dict[str, Any] = {"markdown": "", "tables": [], "text": text,
                                 "error": None, "page_count": 0}
        if ext == ".md":
            parse["markdown"] = text
        ocr_report: dict[str, Any] = {"skipped": "text-based file (no OCR needed)"}
    else:
        parse = parse_document(p)
        text = parse.get("text") or ""
        ocr_report = {}
        # Image files, or non-text binary docs docling couldn't parse → OCR.
        if ext in _IMG_EXTS or (not text.strip() and ext not in _TEXT_EXTS):
            ocr_text, ocr_report = ocr_image(p)
            if len(ocr_text) > len(text):
                text = ocr_text
                parse["text"] = text
        if not text.strip():
            text = _read_fallback(p)
            parse["text"] = text

    entities = extract_entities(text)
    cls = classify(text, entities)                          # content-aware
    typed_fields = extract_typed_fields(text, cls.get("category"), entities)  # per-type schema
    return make_report(str(p), parse, entities, cls, typed_fields, ocr_report)


def _read_fallback(p: Path) -> str:
    """Open xlsx/docx/csv/txt/json/md if docling missed them."""
    try:
        if p.suffix == ".xlsx":
            return "\n".join(
                df.to_string(index=False)
                for df in pd.read_excel(p, sheet_name=None).values()
            )
        if p.suffix == ".docx":
            import docx
            return "\n".join(par.text for par in docx.Document(str(p)).paragraphs)
        if p.suffix in (".csv", ".txt", ".md", ".json", ".xml", ".html", ".htm", ".log"):
            data = p.read_text(encoding="utf-8", errors="ignore")
            # JSON manifests and CSV blobs are readable as-is; MD maps directly.
            return data
    except Exception:
        pass
    return ""


# ---------------------------------------------------------------------------
# CLI entry
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("usage: pipeline.py <file>")
        sys.exit(1)
    report = run_full_pipeline(sys.argv[1])
    print(json.dumps(report, indent=2, default=str))